"""
parser.py
----------
Handles reading raw text out of uploaded resume files (PDF / DOCX).

Beginner notes:
- A PDF resume is parsed using two libraries (PyPDF2 and pdfminer.six).
  We try PyPDF2 first because it's fast. If it fails to find text
  (common with scanned or oddly-formatted PDFs), we fall back to
  pdfminer.six, which is slower but often more accurate.
- A DOCX resume is parsed with python-docx, which reads paragraphs
  and tables directly from the Word XML.
- clean_text() removes extra whitespace/junk characters so downstream
  NLP (SpaCy/NLTK/regex) has consistent input.
"""

import io
import re
import logging
from pathlib import Path
from typing import Union

import PyPDF2
from pdfminer.high_level import extract_text as pdfminer_extract_text
import docx  # python-docx

# Configure a module-level logger so errors are traceable without
# crashing the whole Streamlit app.
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)


def extract_pdf_text(file: Union[str, Path, io.BytesIO]) -> str:
    """
    Extract raw text from a PDF file.

    Args:
        file: A file path (str/Path) OR an in-memory file object
              (e.g. the object returned by Streamlit's file_uploader).

    Returns:
        Extracted text as a single string. Returns "" if extraction
        completely fails (e.g. the file is corrupted).
    """
    text = ""

    # --- Attempt 1: PyPDF2 (fast, works for most "normal" PDFs) ---
    try:
        # PyPDF2 needs the stream position reset to 0 if it's a
        # Streamlit-uploaded file object that may have been read before.
        if hasattr(file, "seek"):
            file.seek(0)

        reader = PyPDF2.PdfReader(file)
        pages_text = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            pages_text.append(page_text)
        text = "\n".join(pages_text).strip()
    except Exception as e:
        logger.warning(f"PyPDF2 failed to extract text: {e}")

    # --- Attempt 2: pdfminer.six fallback if PyPDF2 got little/no text ---
    if len(text) < 30:  # heuristic: almost no text means extraction failed
        try:
            if hasattr(file, "seek"):
                file.seek(0)
            text = pdfminer_extract_text(file) or ""
        except Exception as e:
            logger.error(f"pdfminer.six also failed to extract text: {e}")
            text = text or ""

    return clean_text(text)


def extract_docx_text(file: Union[str, Path, io.BytesIO]) -> str:
    """
    Extract raw text from a DOCX file, including text inside tables
    (many resumes use tables for layout).

    Args:
        file: A file path (str/Path) OR an in-memory file object.

    Returns:
        Extracted text as a single string. Returns "" on failure.
    """
    try:
        if hasattr(file, "seek"):
            file.seek(0)

        document = docx.Document(file)

        parts = []

        # Paragraph text (the normal resume body)
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

        # Table text (skills/experience are sometimes inside tables)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)

        text = "\n".join(parts)
        return clean_text(text)

    except Exception as e:
        logger.error(f"python-docx failed to extract text: {e}")
        return ""


def clean_text(raw_text: str) -> str:
    """
    Normalize extracted text so later steps (regex, SpaCy, NLTK)
    behave consistently.

    Steps:
    1. Replace tabs/newlines with single spaces (but keep line breaks
       for readability where useful).
    2. Collapse multiple spaces into one.
    3. Strip weird non-printable / control characters.
    4. Trim leading/trailing whitespace.
    """
    if not raw_text:
        return ""

    text = raw_text

    # Remove non-printable / control characters (keep standard ASCII + common punctuation)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)

    # Normalize newlines/tabs to spaces but keep paragraph breaks somewhat
    text = re.sub(r"[\t\r]+", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)

    # Collapse multiple spaces
    text = re.sub(r" {2,}", " ", text)

    return text.strip()


def detect_and_extract(file, filename: str) -> str:
    """
    Convenience wrapper: looks at the file extension and calls the
    correct extractor. Used by app.py so it doesn't need to know
    which parser function to call.

    Args:
        file: file-like object or path.
        filename: original filename (used to detect extension).

    Returns:
        Extracted, cleaned text. Empty string if format unsupported
        or extraction fails.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return extract_pdf_text(file)
    elif suffix == ".docx":
        return extract_docx_text(file)
    else:
        logger.warning(f"Unsupported file format: {suffix}")
        return ""
